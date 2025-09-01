import os
import random
from flask import Flask, render_template, request, send_from_directory, flash, redirect, url_for
from docx import Document
from werkzeug.utils import secure_filename

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
ALLOWED_EXTENSIONS = {"txt"}

app = Flask(__name__)
app.secret_key = "secretkey"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["OUTPUT_FOLDER"] = OUTPUT_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Classes et fonctions pour générer QCM (adapté de main.py)
class Genquestionqcm:
    def __init__(self, question, reponses, juste):
        self.question = question
        self.reponses = reponses
        self.juste = juste

def lectureFichier(nomDeFichier):
    listeQuestions = []
    with open(nomDeFichier, 'r', encoding='utf-8') as file:
        lines = file.readlines()
        index = 0
        while index < len(lines):
            line = lines[index].strip()
            if line.endswith("?"):
                question = line
                rep_juste = lines[index+5].strip()
                reponses = [lines[index+1].strip(), lines[index+2].strip(), lines[index+3].strip(), lines[index+4].strip()]
                listeQuestions.append(Genquestionqcm(question, reponses, rep_juste))
                index += 6
            else:
                index += 1
    return listeQuestions

def qcmdocx(liste, sujet_num, folder):
    qcm = Document()
    for x in liste:
        qcm.add_paragraph(x.question)
        for r in x.reponses:
            qcm.add_paragraph(r)
        qcm.add_paragraph()
    path = os.path.join(folder, f"quiz_questions{sujet_num}.docx")
    qcm.save(path)
    return path

def repdocx(reponses, sujet_num, folder):
    rep = Document()
    rep.add_paragraph(reponses)
    path = os.path.join(folder, f"correction{sujet_num}.docx")
    rep.save(path)
    return path

def creation_sujets(liste, nb_suj, nb_qst):
    sujets = []
    reponses_sujets = []
    for i in range(nb_suj):
        sujet = random.sample(liste, nb_qst)
        reponses = "".join([q.juste for q in sujet])
        sujets.append(sujet)
        reponses_sujets.append(reponses)
    return (sujets, reponses_sujets)

# Flask routes
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        file = request.files.get("qcm_file")
        nb_suj = request.form.get("nb_suj")
        nb_qst = request.form.get("nb_qst")
        if not file or not allowed_file(file.filename):
            flash("Veuillez uploader un fichier .txt valide")
            return redirect(request.url)
        try:
            nb_suj = int(nb_suj)
            nb_qst = int(nb_qst)
            if nb_qst % 5 != 0:
                flash("Le nombre de questions doit être un multiple de 5")
                return redirect(request.url)
        except ValueError:
            flash("Veuillez entrer des nombres valides")
            return redirect(request.url)

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(filepath)

        questions = lectureFichier(filepath)
        sujets, reponses = creation_sujets(questions, nb_suj, nb_qst)

        files_generated = []
        for i in range(nb_suj):
            qcm_path = qcmdocx(sujets[i], i+1, app.config["OUTPUT_FOLDER"])
            rep_path = repdocx(reponses[i], i+1, app.config["OUTPUT_FOLDER"])
            files_generated.extend([qcm_path, rep_path])

        return render_template("result.html", files=files_generated)

    return render_template("index.html")

@app.route("/download/<path:filename>")
def download(filename):
    return send_from_directory(app.config["OUTPUT_FOLDER"], filename, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
